#!/usr/bin/env python3
"""Generate Plaza Premium Lounge Website Analysis Report as Word Document."""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(6)

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_header_row(table, row_idx, texts, color="800040"):
    row = table.rows[row_idx]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        set_cell_shading(cell, color)

def add_data_row(table, row_idx, texts):
    row = table.rows[row_idx]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = str(text)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

def add_screenshot(doc, path, caption, width=5.0):
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=Inches(width))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cap.runs:
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(8)
                cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            return True
        except Exception as e:
            doc.add_paragraph(f"[Screenshot: {path} - {str(e)}]")
            return False
    else:
        doc.add_paragraph(f"[Screenshot not available: {path}]")
        return False

# ═══════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════
doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")
title = doc.add_heading("Website Analysis Report", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading("www.plazapremiumlounge.com", level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")
cover_info = doc.add_paragraph("AEM Edge Delivery Services Migration Assessment")
cover_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_info.runs[0].font.size = Pt(16)
cover_info.runs[0].font.color.rgb = RGBColor(128, 0, 64)
doc.add_paragraph("")
date_p = doc.add_paragraph("Date: February 2026")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")
scope_p = doc.add_paragraph("Scope: All pages excluding /en-uk/find/* location pages")
scope_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
scope_p.runs[0].bold = True
doc.add_paragraph("")
platform_p = doc.add_paragraph()
platform_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = platform_p.add_run("Current Platform: Kentico CMS (ASP.NET WebForms)")
run.bold = True
run.font.size = Pt(14)
doc.add_page_break()

# ═══════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Templates Inventory",
    "3. Blocks / Components Catalog",
    "4. Page Counts by Template",
    "5. Integrations Analysis",
    "6. Complex Use Cases & Observations",
    "7. Migration Effort Estimates",
    "Appendix: Screenshots"
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ═══════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "This report provides a comprehensive analysis of www.plazapremiumlounge.com for migration "
    "to Adobe Experience Manager (AEM) Edge Delivery Services. The site is currently built "
    "on Kentico CMS running ASP.NET WebForms and serves as the primary consumer website for "
    "Plaza Premium Lounge, the world's first and largest award-winning independent airport lounge network."
)
doc.add_paragraph(
    "The analysis covers all non-location pages (excluding URLs starting with /en-uk/find/), "
    "which include the homepage, content pages, partner offers, FAQs, legal pages, product pages, "
    "campaign pages, news archive, and form pages. The site supports 3 language variants (en-uk, zh-cn, zh-hk)."
)

doc.add_heading("Key Findings", level=2)
findings = [
    "10 distinct page templates identified with varying complexity levels (Low to High)",
    "42 unique non-location pages in scope (plus ~63 news articles), totaling ~107 pages for migration",
    "29 reusable blocks/components cataloged (7 global + 22 content blocks)",
    "15+ third-party integrations including Netcore Smartech, Google Analytics 4, Facebook Pixel, Pantheon Lab AI chatbot",
    "Complex booking widget on homepage integrates with external booking.plazapremiumlounge.com",
    "3 language versions: English (en-uk), Simplified Chinese (zh-cn), Traditional Chinese (zh-hk)",
    "XML sitemaps are incomplete - only 1 non-location page appears in the en-uk sitemap",
    "~656 location pages (/en-uk/find/*) excluded from scope but documented for reference",
    "Current CMS: Kentico CMS with ASP.NET WebForms, jQuery 3.1.1, Bootstrap, Owl Carousel",
    "Several interactive components require custom EDS block development: booking widget, FAQ accordion, awards timeline, interactive games",
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')
doc.add_page_break()

# ═══════════════════════════════════════════
# 2. TEMPLATES INVENTORY
# ═══════════════════════════════════════════
doc.add_heading("2. Templates Inventory", level=1)
doc.add_paragraph(
    "The following table lists all unique page templates identified across the site. "
    "Each template represents a distinct layout pattern used for one or more pages."
)

templates_data = [
    ["T1", "Homepage", "High",
     "Hero slider with integrated booking widget (city autocomplete, date/time pickers, guest counter), Latest Deals card carousel, New Lounges carousel, promotional banner, Skytrax award badge. Most complex and interactive page.",
     "https://www.plazapremiumlounge.com/en-uk"],
    ["T2", "Content Page", "Medium",
     "Rich editorial content with multiple sections: text blocks, image galleries, service cards, alternating media-text layouts, timeline components. Variants: About Us (company info, services, values, founder bio, awards timeline), PPG Entertainment (hero + alternating service cards), Chef Series (hero + chef grid + PDF download).",
     "https://www.plazapremiumlounge.com/en-uk/discover/about-us\nhttps://www.plazapremiumlounge.com/en-uk/discover/ppg-entertainment\nhttps://www.plazapremiumlounge.com/en-uk/chef-series"],
    ["T3", "Partner Offers Listing", "Medium",
     "2-column card grid displaying 23 promotional offer cards. Each card: thumbnail image, title, description, 'Read more' CTA. Filter toggle at top. No pagination - all offers rendered on single page.",
     "https://www.plazapremiumlounge.com/en-uk/discover/partner-offers"],
    ["T4", "Partner Offer Detail", "Medium",
     "Promotional landing pages with hero banner, structured offer details, CTA buttons. Some variants include tabbed navigation (brand switching), interactive game components, and lounge booking widgets at bottom. 23 pages with variable complexity.",
     "https://www.plazapremiumlounge.com/en-uk/discover/partner-offers/lny2026\nhttps://www.plazapremiumlounge.com/en-uk/discover/partner-offers/pplsinxosc"],
    ["T5", "FAQ Page", "Medium",
     "Two-column layout: left sidebar with 7 category links, right content area with accordion Q&A. Bootstrap accordion with expand/collapse chevrons. 'Expand All' toggle button. Categories: General, Lounge Reservation, Lounge Usage, Lounge Access Entitlement, Gift Card & PPL Pass, Airport Dining, Promotional Code.",
     "https://www.plazapremiumlounge.com/en-uk/discover/faqs/general"],
    ["T6", "Form Page", "High",
     "Complex multi-field forms with cascading dropdowns (country > city), IntlTelInput phone input with country code, date pickers, checkbox groups, CAPTCHA verification, form validation. Contact Us variant adds phone directory + HQ address. Group Bookings variant adds service selection + adult/child counts.",
     "https://www.plazapremiumlounge.com/en-uk/discover/contact-us\nhttps://www.plazapremiumlounge.com/en-uk/discover/group-bookings"],
    ["T7", "Legal / Text Page", "Low",
     "Long-form single-column text with numbered sections and sub-clauses. No images, no interactive elements. Three pages: Terms & Conditions (longest, Part A + Part B), Privacy Policy, Terms of Use.",
     "https://www.plazapremiumlounge.com/en-uk/discover/terms-and-conditions\nhttps://www.plazapremiumlounge.com/en-uk/discover/data-privacy-security-policy\nhttps://www.plazapremiumlounge.com/en-uk/discover/terms-of-use"],
    ["T8", "Product / E-commerce Page", "High",
     "PPL Pass product listing: hero banner with slider, 4-icon benefits strip, search/filter ('I'm traveling to...') with tag-based region filters, 3-column product card grid with pricing/features/CTAs. Smart Traveller promotion section at bottom.",
     "https://www.plazapremiumlounge.com/en-uk/ppl-pass/ppl-pass-home"],
    ["T9", "Campaign / Landing Page", "High",
     "Custom marketing pages with branded hero banners, destination showcase cards with Vimeo video links, chef profile grids with circular portraits, downloadable PDF resources, interactive game elements. Highly bespoke layouts.",
     "https://www.plazapremiumlounge.com/en-US/campaign/yourdestinationbeforedeparture"],
    ["T10", "News Archive", "Medium",
     "Article listing with year/month dropdown filters. 63 articles in vertical list (thumbnail + title + excerpt). No pagination - all articles on single page. Article URLs follow /news/{month-year}/{slug} pattern.",
     "https://www.plazapremiumlounge.com/en-uk/discover/news"],
]

table = doc.add_table(rows=1 + len(templates_data), cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, 0, ["ID", "Template Name", "Complexity", "Description & Reasoning", "Reference URL(s)"])
for i, row_data in enumerate(templates_data):
    add_data_row(table, i + 1, row_data)
for row in table.rows:
    row.cells[0].width = Cm(1.2)
    row.cells[1].width = Cm(2.8)
    row.cells[2].width = Cm(1.8)
    row.cells[3].width = Cm(7)
    row.cells[4].width = Cm(4.5)

doc.add_paragraph("")
doc.add_heading("Template Screenshots", level=2)

template_screenshots = [
    ("/workspace/homepage-full.png", "T1: Homepage - Hero slider with booking widget, promotions carousel, new lounges"),
    ("/workspace/about-us-full.png", "T2: Content Page - About Us (services, founder, awards timeline)"),
    ("/workspace/partner-offers-listing.png", "T3: Partner Offers Listing - 2-column card grid with filter"),
    ("/workspace/partner-offer-detail.png", "T4: Partner Offer Detail - LNY 2026 campaign with tabs"),
    ("/workspace/faqs-general.png", "T5: FAQ Page - Sidebar navigation with accordion Q&A"),
    ("/workspace/contact-us-page.png", "T6: Form Page - Contact Us with enquiry form and phone directory"),
    ("/workspace/terms-conditions.png", "T7: Legal Page - Terms and Conditions"),
    ("/workspace/ppl-pass-page.png", "T8: Product Page - PPL Pass with pricing cards and search"),
    ("/workspace/campaign-page.png", "T9: Campaign Page - Your Destination Before Departure"),
    ("/workspace/news-archive.png", "T10: News Archive - Article listing with date filtering"),
]

for filepath, caption in template_screenshots:
    add_screenshot(doc, filepath, caption, width=4.5)
    doc.add_paragraph("")

doc.add_page_break()

# ═══════════════════════════════════════════
# 3. BLOCKS / COMPONENTS CATALOG
# ═══════════════════════════════════════════
doc.add_heading("3. Blocks / Components Catalog", level=1)
doc.add_paragraph(
    "The following catalog identifies all reusable blocks and components across the site. "
    "Design variations of the same content model are noted rather than treated as separate blocks."
)

doc.add_heading("3.1 Global Components (7 blocks)", level=2)

global_blocks = [
    ["B01", "Global Header / Navigation", "High",
     "Sticky header bar with logo, primary nav (Find Destinations, Lounge Pass, Discover More, Book Now), "
     "utility links (Campaign CTA, Language/Currency selector, Login). Mobile: hamburger menu with full-screen "
     "slide-in panel, language tabs, currency dropdown. Includes mega-menu for destinations with region/country/city hierarchy.",
     "All pages"],
    ["B02", "Global Footer", "Medium",
     "Multi-tier footer: (1) Payment options icons (Visa, MC, Amex, PayPal, Alipay, WeChat Pay), "
     "(2) Newsletter signup form (3 fields + Submit + privacy consent), (3) Social media links (Facebook, Instagram, WeChat), "
     "(4) Plaza Premium Group brand family logos by category (Lounge, Concierge, Hotel, F&B, Rewards), "
     "(5) Copyright bar with Terms/Privacy links.",
     "All pages"],
    ["B03", "Cookie Consent Banner", "Low",
     "Horizontal bar with cookie notice text, privacy policy link, 'I Understand' dismiss button. Powered by ConsenTag (consentag.eu).",
     "All pages"],
    ["B04", "AI Chat Widget", "Medium",
     "Floating AI chatbot button (bottom-right) powered by Pantheon Lab. Opens conversational AI interface in iframe overlay.",
     "All pages"],
    ["B05", "Promotional Popup", "Medium",
     "Modal overlay popup (Netcore Smartech) promoting Smart Traveller membership with 20% off messaging and close button.",
     "Homepage + selected pages"],
    ["B06", "Login / Smart Traveller Modal", "High",
     "Centered modal dialog for Smart Traveller account login. Fields: Email, Password (show/hide toggle), Remember Me. "
     "Actions: Login, Forgot Password, Sign Up Now. Integrates with booking system for session handoff.",
     "All pages (via header)"],
    ["B07", "Breadcrumb Navigation", "Low",
     "Horizontal breadcrumb trail (Home | Category | Page). Pipe-separated links. Appears below header on all interior pages.",
     "All interior pages"],
]

table = doc.add_table(rows=1 + len(global_blocks), cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, 0, ["ID", "Block Name", "Complexity", "Description & Behaviour", "Reference URL(s)"])
for i, row_data in enumerate(global_blocks):
    add_data_row(table, i + 1, row_data)
for row in table.rows:
    row.cells[0].width = Cm(1.2)
    row.cells[1].width = Cm(3)
    row.cells[2].width = Cm(1.8)
    row.cells[3].width = Cm(7.5)
    row.cells[4].width = Cm(3.8)

doc.add_paragraph("")
doc.add_heading("Global Component Screenshots", level=3)
add_screenshot(doc, "/workspace/component-header.png", "B01: Global Header", width=5.0)
doc.add_paragraph("")
add_screenshot(doc, "/workspace/component-footer.png", "B02: Global Footer", width=5.0)
doc.add_paragraph("")
add_screenshot(doc, "/workspace/component-cookie-banner.png", "B03: Cookie Consent Banner", width=5.0)
doc.add_paragraph("")
add_screenshot(doc, "/workspace/login-popup.png", "B06: Login / Smart Traveller Modal", width=3.5)
doc.add_paragraph("")
add_screenshot(doc, "/workspace/mobile-nav.png", "B01 (Mobile): Slide-in Navigation Panel", width=2.5)

doc.add_page_break()

doc.add_heading("3.2 Content Blocks (22 blocks)", level=2)

content_blocks = [
    ["B08", "Hero Slider / Banner", "High",
     "Full-width hero image slider with text overlay. Design variations: (a) Homepage: with integrated booking widget, (b) Campaign: themed imagery with CTA, (c) Content: static banner with text overlay. Uses MasterSlider/Owl Carousel.",
     "Homepage, Campaign, Entertainment, PPL Pass"],
    ["B09", "Booking Widget", "High",
     "Inline search/booking form: City autocomplete, Date picker, Time selector, Guest counter (Adults/Children), 'Book Now' CTA. "
     "Integrates with external booking.plazapremiumlounge.com. Disclaimer text about advance booking.",
     "Homepage (hero section)"],
    ["B10", "Promotional Card Grid", "Medium",
     "Card collection with image, title, description, CTA. Design variations: (a) 4-card horizontal carousel with arrows (homepage), "
     "(b) 2-column responsive grid with filter toggle (offers listing). Same content model.",
     "Homepage, Partner Offers listing"],
    ["B11", "New Lounges Carousel", "Medium",
     "Horizontal carousel of lounge location cards with background images, location name, 'Opened in [Date]' badge. "
     "Pagination dots, 3 cards visible at a time, 13+ items.",
     "Homepage"],
    ["B12", "Service Cards (Alternating)", "Medium",
     "Two-column media+text blocks in alternating layout (image-left/text-right, then reversed). "
     "Each: service logo/image, description, optional external link.",
     "About Us (Our Services), PPG Entertainment"],
    ["B13", "Core Values / Icon Strip", "Low",
     "Horizontal row of 3-4 circular icons with labels. Design variations: (a) 3 icons (Innovation, Team Work, Excellent Service), "
     "(b) 4 icons (Access, Flexibility, Duration, Gift). Same content model.",
     "About Us, PPL Pass"],
    ["B14", "Person / Founder Profile", "Low",
     "Two-column: portrait image (left) + biographical text paragraphs (right). Name heading + multi-paragraph content.",
     "About Us"],
    ["B15", "Awards Timeline", "High",
     "Interactive horizontal timeline (1998-2025) with clickable year markers. Each year reveals awards/milestones panel. "
     "Scrollable horizontal year selector with active state.",
     "About Us"],
    ["B16", "FAQ Accordion", "Medium",
     "Bootstrap accordion with clickable H4 question headings, expandable answer panels, chevron indicators. "
     "'Expand All' toggle. Combined with sidebar category navigation (B17) for FAQ template.",
     "FAQ pages (7 subcategories)"],
    ["B17", "Sidebar Category Navigation", "Low",
     "Vertical nav panel with category links and active state highlighting (maroon). Left sidebar on FAQ pages.",
     "FAQ pages"],
    ["B18", "Contact Form (Complex)", "High",
     "Multi-field form: cascading dropdowns (country > city), IntlTelInput phone input, date pickers, checkbox groups, "
     "CAPTCHA image, privacy consent, form validation. Variants: Contact Us enquiry, Group Bookings request.",
     "Contact Us, Group Bookings"],
    ["B19", "Contact Directory", "Medium",
     "3-column grid of phone numbers (16 countries) + structured dept email contacts + physical HQ address.",
     "Contact Us"],
    ["B20", "Product Card Grid", "High",
     "3-column product cards: pass image, name, pricing, feature bullets, dual CTAs (Buy Now / Learn More). "
     "Includes search input ('I'm traveling to...') + tag-based region/country filters.",
     "PPL Pass page"],
    ["B21", "Tab Navigation", "Medium",
     "Horizontal button bar for switching between content sections. Used for brand-specific offer tabs "
     "(Plaza Premium Lounge, Plaza Premium First, Allways, Aerotel).",
     "LNY 2026 offer page"],
    ["B22", "Destination Showcase Grid", "High",
     "2-column grid of destination cards: city background image, name overlay, expandable detail panel, "
     "'Discover More' link, 'Watch Video' link (Vimeo). 9+ destinations.",
     "Campaign page"],
    ["B23", "Chef Profile Grid", "Medium",
     "Circular portrait photos with chef name + city labels. 4 per row, centered. Clickable. "
     "Accompanied by recipe booklet download section (PDF).",
     "Chef Series page"],
    ["B24", "News Article Card List", "Medium",
     "Vertical list: thumbnail image (left) + title link + excerpt (right). "
     "Year/month dropdown filters. 63 articles, no pagination.",
     "News Archive"],
    ["B25", "Promotional Banner / CTA", "Low",
     "Full-width section: heading, descriptive text, CTA button. 'Save up to 20%' messaging + Smart Traveller promo. "
     "Design variations: (a) text-heavy with bullet points, (b) image+text with CTA.",
     "Homepage, PPL Pass"],
    ["B26", "Skytrax Award Badge", "Low",
     "Full-width image banner displaying Skytrax World Airline Awards certification logos.",
     "Homepage"],
    ["B27", "Newsletter Signup Form", "Low",
     "Inline 3-field form (First Name, Last Name, Email) + Submit + privacy consent. Part of global footer.",
     "All pages (footer)"],
    ["B28", "Interactive Game", "High",
     "'Celestial Sky Race' interactive HTML5 game with 7 destination buttons. Bespoke gamification for LNY campaign.",
     "LNY 2026 offer page"],
    ["B29", "Offer-Specific Lounge Listing", "Medium",
     "Location-specific lounge cards at bottom of offer pages. Shows eligible lounges with rating + 'Apply' button.",
     "Partner offer detail pages"],
]

table = doc.add_table(rows=1 + len(content_blocks), cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, 0, ["ID", "Block Name", "Complexity", "Description & Behaviour", "Reference URL(s)"])
for i, row_data in enumerate(content_blocks):
    add_data_row(table, i + 1, row_data)
for row in table.rows:
    row.cells[0].width = Cm(1.2)
    row.cells[1].width = Cm(3)
    row.cells[2].width = Cm(1.8)
    row.cells[3].width = Cm(7.5)
    row.cells[4].width = Cm(3.8)

doc.add_paragraph("")
doc.add_heading("Content Block Screenshots", level=3)
add_screenshot(doc, "/workspace/component-booking-widget.png", "B08/B09: Hero Slider with Booking Widget", width=5.0)
doc.add_paragraph("")
add_screenshot(doc, "/workspace/component-offer-cards.png", "B10: Promotional Card Grid (offers listing)", width=5.0)
doc.add_paragraph("")
add_screenshot(doc, "/workspace/component-lounge-card.png", "B29: Lounge listing card pattern", width=5.0)

doc.add_page_break()

# ═══════════════════════════════════════════
# 4. PAGE COUNTS BY TEMPLATE
# ═══════════════════════════════════════════
doc.add_heading("4. Page Counts by Template", level=1)
doc.add_paragraph(
    "The following table provides page counts per template type with migration approach assessment. "
    "Location pages (/en-uk/find/*) are excluded from scope but listed for reference."
)

doc.add_heading("4.1 In-Scope Pages (Non-Location)", level=2)

page_counts_data = [
    ["T1", "Homepage", "1", "Manual", "Complex interactive booking widget, carousels, dynamic content. Requires custom block development."],
    ["T2", "Content Page", "3", "Semi-Auto", "About Us, PPG Entertainment, Chef Series. Rich editorial content with custom layouts."],
    ["T3", "Partner Offers Listing", "1", "Semi-Auto", "Card grid with filter. Requires custom listing block."],
    ["T4", "Partner Offer Detail", "23", "Semi-Auto", "23 individual offer pages. Variable complexity. Bulk migration for simpler ones; manual for rich campaign pages."],
    ["T5", "FAQ Page", "7", "Semi-Auto", "7 sub-category pages sharing same template. Accordion needs custom development; content semi-automated."],
    ["T6", "Form Page", "2", "Manual", "Contact Us, Group Bookings. Complex forms with CAPTCHA, cascading dropdowns, phone input."],
    ["T7", "Legal / Text Page", "3", "Automated", "Terms & Conditions, Privacy Policy, Terms of Use. Pure text, straightforward markdown conversion."],
    ["T8", "Product Page (PPL Pass)", "1", "Manual", "E-commerce product listing with pricing, search/filter, buy CTAs. Custom block development needed."],
    ["T9", "Campaign / Landing Page", "2", "Manual", "Your Destination Before Departure + Chef Series. Bespoke designs requiring individual treatment."],
    ["T10", "News Archive", "1 + ~63 articles", "Semi-Auto", "Archive listing + 63 news articles. Listing needs custom block; articles can be bulk migrated."],
    ["", "TOTAL IN-SCOPE", "~107 pages", "", "42 unique template instances + ~63 news articles + redirects"],
]

table = doc.add_table(rows=1 + len(page_counts_data), cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, 0, ["ID", "Template", "Page Count", "Migration Type", "Notes"])
for i, row_data in enumerate(page_counts_data):
    add_data_row(table, i + 1, row_data)
    if row_data[0] == "":
        for j in range(5):
            cell = table.rows[i+1].cells[j]
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
for row in table.rows:
    row.cells[0].width = Cm(1.2)
    row.cells[1].width = Cm(3.2)
    row.cells[2].width = Cm(2.2)
    row.cells[3].width = Cm(2.3)
    row.cells[4].width = Cm(8.5)

doc.add_paragraph("")
doc.add_heading("4.2 Out-of-Scope Location Pages (Reference)", level=2)

loc_data = [
    ["en-uk /find/ pages", "656", "City listing + individual lounge detail pages"],
    ["zh-cn /find/ pages", "631", "Simplified Chinese location pages"],
    ["zh-hk /find/ pages", "767", "Traditional Chinese location pages"],
    ["Total location pages", "2,054", "All language variants combined"],
]

table = doc.add_table(rows=1 + len(loc_data), cols=3)
table.style = 'Table Grid'
add_header_row(table, 0, ["Category", "Count", "Description"])
for i, row_data in enumerate(loc_data):
    add_data_row(table, i + 1, row_data)

doc.add_paragraph("")
doc.add_heading("4.3 Migration Approach Summary", level=2)

migration_summary = [
    ["Automated", "3 pages", "Legal/text pages", "Straightforward text-to-markdown conversion"],
    ["Semi-Automated", "~97 pages", "Offer details (23), FAQs (7), news articles (63), content (3), listing (1)", "Template-based extraction with manual review"],
    ["Manual", "7 pages", "Homepage, form pages (2), product page, campaigns (2), 404 page", "Custom block development, bespoke layouts"],
]

table = doc.add_table(rows=1 + len(migration_summary), cols=4)
table.style = 'Table Grid'
add_header_row(table, 0, ["Approach", "Page Count", "Pages", "Reasoning"])
for i, row_data in enumerate(migration_summary):
    add_data_row(table, i + 1, row_data)
doc.add_page_break()

# ═══════════════════════════════════════════
# 5. INTEGRATIONS ANALYSIS
# ═══════════════════════════════════════════
doc.add_heading("5. Integrations Analysis", level=1)
doc.add_paragraph(
    "The following catalog documents all third-party integrations and embedded services detected across the site."
)

doc.add_heading("5.1 Analytics & Advertising", level=2)

analytics_data = [
    ["Google Tag Manager", "Tag Management", "GTM-5W3S93M", "Low", "Central tag container managing all analytics/marketing tags.", "All pages"],
    ["Google Analytics 4", "Analytics", "G-GPQ3D5T5BY", "Low", "Primary web analytics. Loaded via GTM.", "All pages"],
    ["Google Ads Conversion", "Advertising", "ID: 955808396", "Low", "Google Ads conversion tracking pixel.", "All pages"],
    ["DoubleClick Floodlight", "Advertising (iframe)", "Source: 8141010", "Low", "Google DV360 campaign attribution.", "All pages"],
    ["Bing UET", "Advertising", "Tag: 343126860", "Low", "Microsoft/Bing Ads conversion tracking.", "All pages"],
    ["Yahoo Japan yTag", "Advertising", "Conversion pixel", "Low", "Yahoo Japan advertising tracking.", "All pages"],
    ["Facebook/Meta Pixel", "Advertising", "Pixel: 517402188726028", "Low", "Facebook conversion tracking and audience building.", "All pages"],
]

table = doc.add_table(rows=1 + len(analytics_data), cols=6)
table.style = 'Table Grid'
add_header_row(table, 0, ["Integration", "Type", "Identifier", "Complexity", "Description", "Pages"])
for i, row_data in enumerate(analytics_data):
    add_data_row(table, i + 1, row_data)

doc.add_paragraph("")
doc.add_heading("5.2 Marketing Automation & Engagement", level=2)

marketing_data = [
    ["Netcore Smartech", "Marketing Automation", "Custom SDK", "High",
     "Primary marketing automation: newsletter forms, web push notifications, behavioral popups, session tracking, web personalization. Major integration effort.",
     "All pages"],
    ["Netcore Hansel", "A/B Testing", "SDK v8.4.0", "Medium",
     "Product experience platform for nudges and A/B testing. Part of Netcore suite.",
     "All pages"],
    ["Boxx.ai (Netcore)", "AI Personalization", "API", "Medium",
     "AI-powered personalization/recommendation engine integrated with Smartech.",
     "All pages"],
    ["Pantheon Lab AI Chatbot", "Customer Support", "Iframe widget", "Medium",
     "AI chatbot displayed as floating button. Opens conversational interface in iframe.",
     "All pages"],
    ["ConsenTag", "Consent Management", "Container: 67407556", "Medium",
     "Cookie consent management (consentag.eu). Handles GDPR consent. May need EDS-compatible replacement.",
     "All pages"],
]

table = doc.add_table(rows=1 + len(marketing_data), cols=6)
table.style = 'Table Grid'
add_header_row(table, 0, ["Integration", "Type", "Identifier", "Complexity", "Description", "Pages"])
for i, row_data in enumerate(marketing_data):
    add_data_row(table, i + 1, row_data)

doc.add_paragraph("")
doc.add_heading("5.3 Affiliate & Retargeting", level=2)

affiliate_data = [
    ["Affilired", "Affiliate Tracking", "Merchant: 4426", "Low", "Affiliate commission tracking.", "All pages"],
    ["Denomatic", "Retargeting", "Script ID: 144", "Low", "Display/retargeting ads.", "All pages"],
    ["Sojern", "Travel Advertising", "Pixel: 242598", "Low", "Travel industry targeting.", "Offer pages"],
    ["Red Intelligence", "Retargeting (iframe)", "Account: 68469", "Low", "Retargeting pixel.", "Homepage"],
    ["RTG Ads", "Retargeting", "Ad: 842", "Low", "Retargeting/conversion pixel.", "Homepage"],
    ["Cookieless Data", "Audience Targeting", "PA: 32114", "Low", "Cookieless lookalike targeting.", "Homepage"],
]

table = doc.add_table(rows=1 + len(affiliate_data), cols=6)
table.style = 'Table Grid'
add_header_row(table, 0, ["Integration", "Type", "Identifier", "Complexity", "Description", "Pages"])
for i, row_data in enumerate(affiliate_data):
    add_data_row(table, i + 1, row_data)

doc.add_paragraph("")
doc.add_heading("5.4 External Systems & Booking", level=2)

external_data = [
    ["Booking Engine", "External Application", "booking.plazapremiumlounge.com", "High",
     "External booking system for lounge reservations and PPL Pass purchases. Homepage widget redirects with parameters. Critical integration.",
     "Homepage, PPL Pass, Lounge pages, Offer pages"],
    ["Partner Portal", "External Application", "partnerportal.plaza-network.com", "Low",
     "External partner portal. Currently 404. Simple redirect needed.",
     "Partner Portal page"],
    ["Smart Traveller", "External System", "mysmarttraveller.com", "Medium",
     "Loyalty/rewards platform. Login integration in header modal. External registration/account management.",
     "Login modal, footer, CTAs"],
    ["Vimeo", "Video Hosting", "Embed/Links", "Low",
     "Video hosting for campaign content. Used as linked URLs.",
     "Campaign pages"],
]

table = doc.add_table(rows=1 + len(external_data), cols=6)
table.style = 'Table Grid'
add_header_row(table, 0, ["Integration", "Type", "Identifier", "Complexity", "Description", "Pages"])
for i, row_data in enumerate(external_data):
    add_data_row(table, i + 1, row_data)

doc.add_paragraph("")
doc.add_heading("5.5 Frontend Libraries (to be replaced/removed)", level=2)

libs_data = [
    ["jQuery 3.1.1", "JS Library", "CDN", "Medium", "Core JS library used extensively. Must rewrite in vanilla JS for EDS."],
    ["Bootstrap", "CSS/JS Framework", "Kentico-bundled", "Medium", "Accordion, modals, grid. Replace with custom CSS."],
    ["GSAP (TweenMax) 1.18.0", "Animation", "CDN", "Low", "Scroll-based animations. Replace with CSS animations."],
    ["Owl Carousel / Master Slider", "Carousel", "CMS-bundled", "Medium", "Hero slider, content carousels. Requires custom carousel block."],
    ["Fancybox 3.1.20", "Lightbox", "CDN", "Low", "Image lightbox. Replace with simple EDS lightbox."],
    ["Google Fonts (Poppins, Roboto)", "Typography", "CDN", "Low", "Port to EDS fonts.css."],
    ["Adobe Typekit (inf6ekj)", "Typography", "CDN", "Low", "Load in EDS head.html."],
]

table = doc.add_table(rows=1 + len(libs_data), cols=5)
table.style = 'Table Grid'
add_header_row(table, 0, ["Library", "Type", "Source", "Migration Complexity", "Notes"])
for i, row_data in enumerate(libs_data):
    add_data_row(table, i + 1, row_data)

doc.add_page_break()

# ═══════════════════════════════════════════
# 6. COMPLEX USE CASES
# ═══════════════════════════════════════════
doc.add_heading("6. Complex Use Cases & Observations", level=1)
doc.add_paragraph(
    "The following identifies complex behaviours and edge cases requiring special attention during migration."
)

complex_data = [
    ["CU01", "Booking Widget with City Autocomplete", "1 instance",
     "Homepage",
     "City autocomplete queries API for airport/city matches. Combined with date/time/guest inputs and 'Book Now' redirect to external booking engine with URL parameters. Primary conversion funnel - requires custom EDS block with API connectivity."],
    ["CU02", "Multi-Language Site (3 locales)", "3 variants",
     "All pages (en-uk, zh-cn, zh-hk)",
     "3 language variants with separate URL prefixes, sitemaps, and currency selector. Partial content parity. EDS requires careful URL structure planning and content duplication strategy."],
    ["CU03", "Cascading Dropdown Forms + CAPTCHA", "2 instances",
     "Contact Us, Group Bookings",
     "Country > City cascading dropdowns, IntlTelInput phone input, custom image CAPTCHA (not reCAPTCHA). ASP.NET postback form submission must be redesigned for client-side API calls."],
    ["CU04", "Interactive Campaign Game", "1 instance",
     "LNY 2026 offer page",
     "'Celestial Sky Race' interactive HTML5 game with destination buttons. Bespoke gamification tied to promotional tracking. Must rebuild as custom EDS block or embed as iframe."],
    ["CU05", "Smart Traveller Authentication", "1 instance (all pages)",
     "Header login modal",
     "Login modal connects to membership system. Session management, personalized pricing for members, booking form pre-fill, points display. EDS is static-first; requires client-side auth integration."],
    ["CU06", "No-Pagination Content Loading", "3 instances",
     "News (63 articles), Offers (23 cards), FAQ pages",
     "All content loaded on single page without pagination. News page has 63 items. Affects performance. EDS should implement pagination or lazy loading."],
    ["CU07", "External Booking Engine Deep Links", "5+ touchpoints",
     "Homepage, PPL Pass, Lounge pages, Offer pages",
     "Multiple CTAs redirect to booking.plazapremiumlounge.com with URL parameters (lounge ID, date, guests). Different flows for lounges vs passes. Must maintain all deep link parameters."],
    ["CU08", "Netcore Smartech Deep Integration", "All pages",
     "Site-wide",
     "Newsletter forms, web push, behavioral popups, cross-tab communication, AI personalization (Boxx.ai). Multiple script entry points must be re-integrated in EDS delayed.js without impacting performance."],
    ["CU09", "Awards Timeline (27 years)", "1 instance",
     "About Us page",
     "Interactive horizontal scrollable timeline 1998-2025 with 27 clickable year markers. Content switching on click. Requires custom EDS block with horizontal scroll interaction."],
    ["CU10", "Redirects & Dead Links", "4 instances",
     "Various pages",
     "/faqs > /faqs/general, /ppl-pass-experience > /ppl-pass/ppl-pass-home, /partner-portal > external 404, /about-aerotel/smart-traveller > 404. Redirect rules must be maintained; dead links need cleanup."],
]

table = doc.add_table(rows=1 + len(complex_data), cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, 0, ["ID", "Use Case", "Instances", "Where Found", "Why It's Complex"])
for i, row_data in enumerate(complex_data):
    add_data_row(table, i + 1, row_data)
for row in table.rows:
    row.cells[0].width = Cm(1.2)
    row.cells[1].width = Cm(3)
    row.cells[2].width = Cm(2)
    row.cells[3].width = Cm(3)
    row.cells[4].width = Cm(8)

doc.add_page_break()

# ═══════════════════════════════════════════
# 7. MIGRATION ESTIMATES
# ═══════════════════════════════════════════
doc.add_heading("7. Migration Effort Estimates", level=1)
doc.add_paragraph(
    "The following estimates are based on the analysis of 42 unique in-scope pages (plus ~63 news articles), "
    "29 identified blocks/components, and 15+ third-party integrations. Estimates assume a team experienced with "
    "Edge Delivery Services."
)

doc.add_heading("7.1 Effort Breakdown by Work Stream", level=2)

effort_data = [
    ["WS1", "Project Setup & Architecture", "EDS scaffolding, Git repo, aem.live config, i18n URL structure for 3 languages, CI/CD.", "3-5 days", "-"],
    ["WS2", "Design System Migration", "Extract tokens (colors, typography, spacing), CSS custom properties, fonts.css (Poppins, Roboto, Typekit), responsive breakpoints.", "5-7 days", "-"],
    ["WS3", "Global Components (B01-B07)", "Header/Nav (desktop + mobile mega-menu), Footer, Cookie Consent, Chat Widget embed, Login Modal stub, Breadcrumb, Promo Popup.", "8-12 days", "-"],
    ["WS4", "Blocks - Low Complexity (7 blocks)", "Core Values/Icons (B13), Person Profile (B14), Sidebar Nav (B17), Promo Banner (B25), Award Badge (B26), Newsletter (B27), Breadcrumb.", "4-5 days", "-"],
    ["WS5", "Blocks - Medium Complexity (9 blocks)", "Promo Card Grid (B10), Lounge Carousel (B11), Service Cards (B12), FAQ Accordion (B16), Contact Directory (B19), Tab Nav (B21), Chef Grid (B23), News List (B24), Offer Lounges (B29).", "10-14 days", "-"],
    ["WS6", "Blocks - High Complexity (6 blocks)", "Hero Slider (B08), Booking Widget (B09), Contact Form (B18), Product Card Grid (B20), Destination Showcase (B22), Awards Timeline (B15), Interactive Game (B28).", "14-20 days", "-"],
    ["WS7", "Template Development", "10 page templates (T1-T10) with proper block composition and responsive layouts.", "5-7 days", "-"],
    ["WS8", "Content Migration - Automated", "3 legal/text pages. Direct markdown conversion.", "-", "1-2 days"],
    ["WS9", "Content Migration - Semi-Auto", "97 pages: offers (23), FAQs (7), news articles (63), content (3), listing (1). Template-based extraction + manual review.", "-", "10-15 days"],
    ["WS10", "Content Migration - Manual", "7 pages: Homepage, forms (2), PPL Pass, campaigns (2), 404. Complex interactive layouts.", "-", "8-12 days"],
    ["WS11", "Integration Setup", "GTM + all analytics/ad pixels, Netcore Smartech, Pantheon Lab chatbot, ConsenTag, booking deep links, Smart Traveller auth.", "8-12 days", "-"],
    ["WS12", "QA & Testing", "Cross-browser, responsive, accessibility (WCAG 2.1 AA), performance (Lighthouse), functional testing, content review, link validation, redirects.", "-", "10-15 days"],
    ["WS13", "UAT & Launch", "User acceptance testing, redirect mapping, CDN config, go-live checklist, monitoring.", "-", "5-7 days"],
]

table = doc.add_table(rows=1 + len(effort_data), cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, 0, ["ID", "Work Stream", "Description", "Dev Effort", "Migration/QA"])
for i, row_data in enumerate(effort_data):
    add_data_row(table, i + 1, row_data)
for row in table.rows:
    row.cells[0].width = Cm(1)
    row.cells[1].width = Cm(3)
    row.cells[2].width = Cm(7)
    row.cells[3].width = Cm(2.5)
    row.cells[4].width = Cm(3)

doc.add_paragraph("")
doc.add_heading("7.2 Overall Estimates Summary", level=2)

summary_effort = [
    ["Development (Blocks + Templates)", "49-70 person-days", "WS1-WS7, WS11"],
    ["Content Migration (all approaches)", "19-29 person-days", "WS8-WS10"],
    ["QA & Testing", "10-15 person-days", "WS12"],
    ["UAT & Launch", "5-7 person-days", "WS13"],
    ["TOTAL (Sequential)", "83-121 person-days", "All work streams"],
    ["TOTAL (Parallel Execution)", "55-75 person-days", "With 2-3 devs + 1 QA + 1 content author"],
]

table = doc.add_table(rows=1 + len(summary_effort), cols=3)
table.style = 'Table Grid'
add_header_row(table, 0, ["Category", "Effort Estimate", "Work Streams"])
for i, row_data in enumerate(summary_effort):
    add_data_row(table, i + 1, row_data)
    if "TOTAL" in row_data[0]:
        for j in range(3):
            cell = table.rows[i+1].cells[j]
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

doc.add_paragraph("")
doc.add_heading("7.3 Recommended Team Composition", level=2)

team_data = [
    ["EDS Lead Developer", "1", "Architecture, complex blocks, integration lead"],
    ["EDS Developer", "1-2", "Block development, styling, responsive implementation"],
    ["Content Author / Migration Specialist", "1-2", "Content extraction, markdown authoring, review"],
    ["QA Engineer", "1", "Cross-browser testing, accessibility, performance"],
    ["Project Manager", "1", "Coordination, stakeholder communication, timeline"],
]

table = doc.add_table(rows=1 + len(team_data), cols=3)
table.style = 'Table Grid'
add_header_row(table, 0, ["Role", "Count", "Responsibilities"])
for i, row_data in enumerate(team_data):
    add_data_row(table, i + 1, row_data)

doc.add_paragraph("")
doc.add_heading("7.4 Key Assumptions & Risks", level=2)

risks = [
    "Estimates assume migration of en-uk language ONLY. Additional languages (zh-cn, zh-hk) add ~30-40% effort for content migration and QA.",
    "Location pages (/find/*) are excluded. If brought into scope, 656+ pages need separate bulk migration workstream.",
    "Booking engine integration assumes existing API is accessible and documented. If redesign needed, add significant effort.",
    "Smart Traveller login requires coordination with membership platform team.",
    "Campaign pages (T9) are bespoke - future campaigns will each need individual development.",
    "CAPTCHA on forms requires alternative solution (reCAPTCHA / Turnstile) since EDS is static.",
    "Netcore Smartech re-integration complexity depends on required feature depth (newsletters, push, popups, personalization).",
    "Current jQuery-dependent code must be fully rewritten in vanilla JavaScript for EDS.",
    "Performance target is Lighthouse 100 score.",
    "News article count (63 from 2016-2023) may grow if archive continues.",
]

for r in risks:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading("7.5 Recommended Migration Phases", level=2)

phases = [
    ["Phase 1: Foundation", "Weeks 1-3", "Project setup, design system extraction, global components (header, footer, consent, chat)."],
    ["Phase 2: Core Blocks", "Weeks 3-8", "All block development. Parallel low/medium blocks while high-complexity blocks are architected."],
    ["Phase 3: Templates & Integration", "Weeks 6-10", "Template assembly, integration setup, booking engine connectivity, auth integration."],
    ["Phase 4: Content Migration", "Weeks 8-12", "Automated migration scripting, followed by semi-automated and manual page migration."],
    ["Phase 5: QA & Launch", "Weeks 10-14", "Testing, accessibility audit, performance optimization, UAT, redirect mapping, staged rollout."],
]

table = doc.add_table(rows=1 + len(phases), cols=3)
table.style = 'Table Grid'
add_header_row(table, 0, ["Phase", "Timeline", "Description"])
for i, row_data in enumerate(phases):
    add_data_row(table, i + 1, row_data)

doc.add_page_break()

# ═══════════════════════════════════════════
# APPENDIX: ADDITIONAL SCREENSHOTS
# ═══════════════════════════════════════════
doc.add_heading("Appendix: Additional Screenshots", level=1)

appendix_screenshots = [
    ("/workspace/partner-offer-detail2.png", "Partner Offer Detail - Old Seng Choong Singapore"),
    ("/workspace/group-bookings.png", "Group Bookings Form Page"),
    ("/workspace/entertainment-page.png", "PPG Entertainment Page"),
    ("/workspace/privacy-policy.png", "Privacy Policy Page"),
    ("/workspace/chef-series.png", "Chef Series Campaign Page"),
    ("/workspace/faq-lounge-reservation.png", "FAQ - Lounge Reservation subcategory"),
    ("/workspace/location-city-london.png", "Location City Page - London (out of scope, reference)"),
    ("/workspace/location-lounge-heathrow.png", "Location Lounge Detail - Heathrow T3 (out of scope, reference)"),
    ("/workspace/location-lounge-vancouver.png", "Location Lounge Detail - Vancouver (out of scope, reference)"),
]

for filepath, caption in appendix_screenshots:
    add_screenshot(doc, filepath, caption, width=4.5)
    doc.add_paragraph("")

# ── Save ──
output_path = "/workspace/Plaza_Premium_Lounge_Migration_Analysis.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024 / 1024:.1f} MB")
